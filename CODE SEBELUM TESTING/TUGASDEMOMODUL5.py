from docx import Document
from docx.shared import Inches
import os


filename = 'pesanan_restoran.docx'

def initialize_file():
    if not os.path.exists(filename):
        doc = Document()
        doc.add_heading('Data Pesanan Restoran', level=1)
        doc.add_table(rows=1, cols=4).style = 'Table Grid'
        doc.tables[0].rows[0].cells[0].text = "Nama Pelanggan"
        doc.tables[0].rows[0].cells[1].text = "Menu"
        doc.tables[0].rows[0].cells[2].text = "Status"
        doc.tables[0].rows[0].cells[3].text = "Path Gambar"
        doc.save(filename)

def buat_pesanan(nama, menu, path_gambar):
    doc = Document(filename)
    table = doc.tables[0]
    
  
    for row in table.rows:
        if row.cells[0].text == nama:
            print(f"Pesanan atas nama {nama} sudah ada!")
            return


    new_row = table.add_row()
    new_row.cells[0].text = nama
    new_row.cells[1].text = menu
    new_row.cells[2].text = "Diproses"

    if os.path.exists(path_gambar):  
        paragraph = new_row.cells[3].add_paragraph()  
        run = paragraph.add_run()
        run.add_picture(path_gambar, width=Inches(1))  
    else:
        new_row.cells[3].text = "Gambar tidak ditemukan"


    doc.save(filename)
    print(f"Pesanan untuk {nama} berhasil dibuat.")


def baca_pesanan():
    doc = Document(filename)
    table = doc.tables[0]
    for row in table.rows[1:]:
        print(f"Nama: {row.cells[0].text}, Menu: {row.cells[1].text}, Status: {row.cells[2].text}, Gambar: {row.cells[3].text}")

def perbarui_pesanan(nama, status_baru):
    doc = Document(filename)
    table = doc.tables[0]
    for row in table.rows[1:]:
        if row.cells[0].text == nama:
            row.cells[2].text = status_baru
            doc.save(filename)
            print(f"Status pesanan untuk {nama} berhasil diperbarui menjadi {status_baru}.")
            return
    print(f"Pesanan atas nama {nama} tidak ditemukan.")

def hapus_pesanan(nama):
    doc = Document(filename)
    table = doc.tables[0]
    header_cells = [cell.text for cell in table.rows[0].cells]  
    new_table = doc.add_table(rows=1, cols=4)  
    new_table.style = table.style
    for i, text in enumerate(header_cells):
        new_table.rows[0].cells[i].text = text
    for row in table.rows[1:]:
        if not (row.cells[0].text == nama and row.cells[2].text == "Batal"):
            new_row = new_table.add_row()
            for i, cell in enumerate(row.cells):
                new_row.cells[i].text = cell.text

    table._element.getparent().remove(table._element)
    doc.save(filename)
    print(f"Pesanan atas nama {nama} berhasil dihapus (jika statusnya 'Batal').")

def cari_pesanan(nama):
    doc = Document(filename)
    table = doc.tables[0]
    for row in table.rows[1:]:
        if row.cells[0].text == nama:
            print(f"Pesanan ditemukan - Nama: {row.cells[0].text}, Menu: {row.cells[1].text}, Status: {row.cells[2].text}, Gambar: {row.cells[3].text}")
            return
    print(f"Pesanan atas nama {nama} tidak ditemukan.")


initialize_file()

while True:
    print("\n=== Manajemen Pesanan Restoran ===")
    print("1. Tambah Pesanan")
    print("2. Tampilkan Pesanan")
    print("3. Update Pesanan")
    print("4. Hapus Pesanan")
    print("5. Cari Pesanan")
    print("6. Keluar")
    pilihan = input("Pilih menu: ")

    if pilihan == "1":
        nama = input("Masukkan nama pelanggan: ")
        menu = input("Masukkan menu yang dipesan: ")
        path_gambar = input("Masukkan path gambar: ")
        buat_pesanan(nama, menu, path_gambar)
    elif pilihan == "2":
        baca_pesanan()
    elif pilihan == "3":
        nama = input("Masukkan nama pelanggan yang ingin diupdate: ")
        status_baru = input("Masukkan status baru (Diproses/Selesai/Batal): ")
        perbarui_pesanan(nama, status_baru)
    elif pilihan == "4":
        nama = input("Masukkan nama pelanggan yang ingin dihapus: ")
        hapus_pesanan(nama)
    elif pilihan == "5":
        nama = input("Masukkan nama pelanggan yang ingin dicari: ")
        cari_pesanan(nama)
    elif pilihan == "6":
        print("Program selesai.")
        break
    else:
        print("Pilihan tidak valid, coba lagi.")