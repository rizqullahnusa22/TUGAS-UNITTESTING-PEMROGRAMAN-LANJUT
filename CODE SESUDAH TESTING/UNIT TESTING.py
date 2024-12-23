import unittest
from unittest.mock import patch, MagicMock
from docx import Document
import os


from TUGASDEMOMODUL5 import initialize_file, buat_pesanan, baca_pesanan, perbarui_pesanan, hapus_pesanan, cari_pesanan

class TestPesananRestoran(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        # Setup file untuk testing
        cls.test_filename = 'test_pesanan_restoran.docx'
        if os.path.exists(cls.test_filename):
            os.remove(cls.test_filename)  # Hapus jika ada file sebelumnya
        initialize_file()  # Buat file baru untuk testing

    @classmethod
    def tearDownClass(cls):
        # Hapus file setelah testing selesai
        if os.path.exists(cls.test_filename):
            os.remove(cls.test_filename)

    def test_buat_pesanan(self):
        # Test buat pesanan baru
        buat_pesanan("Rizqullah Atsir Dafa", "Nasi Goreng", "gambar_nasi goreng.jpg")

        # Periksa apakah pesanan berhasil ditambahkan
        doc = Document(self.test_filename)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(table.rows[1].cells[0].text, "Rizqullah Atsir Dafa")
        self.assertEqual(table.rows[1].cells[1].text, "Nasi Goreng")
        self.assertEqual(table.rows[1].cells[2].text, "Diproses")
        self.assertEqual(table.rows[1].cells[3].text, "Gambar Nasi Goreng ditemukan")

    def test_buat_pesanan_duplikat(self):
        # Test untuk memastikan pesanan yang sama tidak ditambahkan dua kal
        buat_pesanan("Rizqullah Atsir Dafa", "Nasi Goreng", "gambar_nasi goreng.jpg")
        with patch('builtins.print') as mocked_print:
            buat_pesanan("Rizqullah Atsir Dafa", "Nasi Goreng", "gambar_nasi goreng.jpg")
            mocked_print.assert_called_with("Pesanan atas nama Rizqullah Atsir Dafa sudah ada!")

    def test_perbarui_pesanan(self):
        # Test untuk memperbarui status pesanan
        buat_pesanan("Rizqullah Atsir Dafa", "Nasi Goreng", "gambar_nasi goreng.jpg")
        perbarui_pesanan("Rizqullah Atsir Dafa", "Selesai")

        # Periksa status baru
        doc = Document(self.test_filename)
        table = doc.tables[0]
        self.assertEqual(table.rows[1].cells[2].text, "Selesai")

    def test_hapus_pesanan(self):
        # Test untuk menghapus pesanan
        buat_pesanan("Rizqullah Atsir Dafa", "Nasi Goreng", "gambar_nasi goreng.jpg")
        hapus_pesanan("Rizqullah Atsir Dafa")

        # Periksa apakah pesanan sudah terhapus
        doc = Document(self.test_filename)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 1)  # Hanya header yang tersisa

    def test_cari_pesanan(self):
        # Test untuk mencari pesanan
        buat_pesanan("Rizqullah Atsir Dafa", "Nasi Goreng", "gambar_nasi goreng.jpg")
        with patch('builtins.print') as mocked_print:
            cari_pesanan("Rizqullah Atsir Dafa")
            mocked_print.assert_called_with("Pesanan ditemukan - Nama: Rizqullah Atsir Dafa, Menu: Nasi Goreng, Status: Diproses, Gambar: Gambar Nasi Goreng ditemukan")

    def test_cari_pesanan_tidak_ada(self):
        # Test untuk mencari pesanan yang tidak ada
        with patch('builtins.print') as mocked_print:
            cari_pesanan("Nonexistent User")
            mocked_print.assert_called_with("Pesanan atas nama Nonexistent User tidak ditemukan.")

if __name__ == '__main__':
    unittest.main()
